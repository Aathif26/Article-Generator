from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate

import streamlit as st
import os
from docx import Document
from docx.shared import Inches
import io

from PIL import Image
import requests

import json
from streamlit_lottie import st_lottie

#load the model
def load_llm(max_tokens, prompt_template):
    #load the locally downloaded model here
    llm = CTransformers(
        model = "llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type="llama",
        max_new_tokens = max_tokens,
        temperature = 0.7
    )
    llm_chain = LLMChain(
        llm= llm,
        prompt= PromptTemplate.from_template(prompt_template)
    )
    print(llm_chain)
    return llm_chain


def fetch_photo(query):
    api_key = '9AYStMxDKxskLhwQO4uCAUsbvewCFaNtaaA8goci4mp6frRqsFSmFdqs'

    url = 'https://api.pexels.com/v1/search'

    headers = {
        'Authorization': api_key
    }

    params = {
        'query': query,
        'per_page': 1
    }

    response = requests.get(url, headers=headers, params=params)

    #Chech if the request was successful
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            src_original_url = photos[0]['src']['original']
            return src_original_url
        else:
            st.write("No photos found for the query")
    
    else:
        st.write(f"Error: {response.status_code}, {response.text}")

    return None


# #Example usage of this Function
# query = "Circular Economy"

# src_original_url = fetch_photo(query)

# if src_original_url:
#     print(f"Original URL for: '{query}' : {src_original_url}")

def create_word_docx(user_input, paragraph, image_input):
    doc = Document()

    doc.add_heading(user_input, level = 1)
    doc.add_paragraph(paragraph)

    doc.add_heading('Image', level=1)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_input.seek(0)
    doc.add_picture(image_stream, width=Inches(4))

    return doc

def load_lottieurl(url):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

img_icon = Image.open('article.png')
# st.beta_set_page_config(page_title="Article Generator")
st.set_page_config(page_title="Article Generator",layout="wide", page_icon=img_icon)


def main():
    #Background Image
    page_bg_img = """
    <style>
    [data-testid = "stAppViewContainer"] {
    background-image: url("https://unsplash.com/photos/NVRRZ5pxX4Q");
    background-size: cover;
    }
    </style>
    """

    # Minimalize the Default Feature
    hide_menu_style = """
        <style>
        /*  MainMenu {visibility: hidden; } */
        footer {visibility: hidden; }
        </style>
        """
    

    st.markdown(hide_menu_style, unsafe_allow_html=True)
    st.markdown(page_bg_img, unsafe_allow_html=True)

    lottie_coding = load_lottieurl("https://lottie.host/7a3a92bf-3e51-4981-8746-90670d1bfeb9/WH6jdZUAXV.json")

    st.title("Article Generator")
    st_lottie(lottie_coding, speed=1,reverse=False, loop=True, quality="high", height=200, width=200)
    user_input = st.text_input("Please Enter the topic or Idea for generating the article")
    image_input = st.text_input("Please Enter the idea for generating the image for the article")

    if len(user_input) and len(image_input) > 0:
        col1, col2, col3 = st.columns([1,2,1])

        with col1:
            st.subheader("Generated Content:")
            st.write("The idea for Article Generation: "+ user_input)
            st.write("The idea of image for Article Generation: "+ image_input)

            prompt_template = """ You are a Software Developer and your task is to write an article on the given topic: {user_input}. 
                                The article must be under 1000 words and should be lengthy."""
            
            llm_call = load_llm(max_tokens=1000, prompt_template = prompt_template)
            result = llm_call(user_input)

            if len(result) > 0:
                st.success("Your Article has been Generated Successfully")
                st.write(result)
            
            else:
                st.error("Sorry Under this topic article are not generated.")

        with col2:
            st.subheader("Fetched Image")
            image_url = fetch_photo(image_input)
            st.image(image_url)

        with col3:
            st.subheader("Download Your Generated Article")
            # image_input = "temp_image.jpg"
            image_response = requests.get(image_url)
            img = Image.open(io.BytesIO(image_response.content))
            doc = create_word_docx(user_input, result['text'], img)

            #Save the word doc to a BytesIO Buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Specify content type and length
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            content_length = len(doc_buffer.getvalue())

            #streamlit download button
            st.download_button(
                label="Download article",
                data = doc_buffer,
                file_name=f"{user_input}.docx",
                key=f"download_{user_input}",
                mime=content_type,
                on_click=None,
                kwargs={'length': content_length},
            )

if __name__ == "__main__":
    main()
    